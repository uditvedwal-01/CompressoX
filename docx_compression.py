import os
from docx.api import Document
from docx.shared import Inches
from PIL import Image
import io
import zipfile
import shutil
import xml.etree.ElementTree as ET
import re
from collections import Counter
import lzma
import bz2
import zlib
import heapq

def apply_lossy_algorithm_1(doc):
    """Aggressive Image Compression: Reduces image quality and dimensions significantly"""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run._element.xpath('.//w:drawing'):
                drawing = run._element.xpath('.//w:drawing')[0]
                blip = drawing.xpath('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                if blip:
                    embed = blip[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed:
                        image_part = doc.part.related_parts[embed]
                        try:
                            img = Image.open(io.BytesIO(image_part.blob))
                            # Reduce dimensions by 70%
                            new_size = (int(img.size[0] * 0.3), int(img.size[1] * 0.3))
                            img = img.resize(new_size, Image.Resampling.LANCZOS)
                            # Convert to grayscale
                            img = img.convert('L')
                            # Save with very low quality
                            buffer = io.BytesIO()
                            img.save(buffer, format='JPEG', quality=20, optimize=True)
                            image_part.blob = buffer.getvalue()
                        except Exception as e:
                            print(f"Error in aggressive compression: {str(e)}")

def apply_lossy_algorithm_2(doc):
    """Smart Image Compression: Adaptive compression based on image content"""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run._element.xpath('.//w:drawing'):
                drawing = run._element.xpath('.//w:drawing')[0]
                blip = drawing.xpath('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                if blip:
                    embed = blip[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed:
                        image_part = doc.part.related_parts[embed]
                        try:
                            img = Image.open(io.BytesIO(image_part.blob))
                            # Calculate compression based on image size
                            if img.size[0] * img.size[1] > 1000000:  # Large image
                                new_size = (int(img.size[0] * 0.5), int(img.size[1] * 0.5))
                                quality = 40
                            else:  # Small image
                                new_size = (int(img.size[0] * 0.7), int(img.size[1] * 0.7))
                                quality = 60
                            img = img.resize(new_size, Image.Resampling.LANCZOS)
                            # Reduce colors for large images
                            if img.size[0] * img.size[1] > 1000000:
                                img = img.convert('P', palette=Image.ADAPTIVE, colors=64)
                            buffer = io.BytesIO()
                            img.save(buffer, format='JPEG', quality=quality, optimize=True)
                            image_part.blob = buffer.getvalue()
                        except Exception as e:
                            print(f"Error in smart compression: {str(e)}")

def apply_lossy_algorithm_3(doc):
    """Color Space Optimization: Focuses on color reduction and DPI adjustment"""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run._element.xpath('.//w:drawing'):
                drawing = run._element.xpath('.//w:drawing')[0]
                blip = drawing.xpath('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                if blip:
                    embed = blip[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed:
                        image_part = doc.part.related_parts[embed]
                        try:
                            img = Image.open(io.BytesIO(image_part.blob))
                            # Set DPI to 72
                            img.info['dpi'] = (72, 72)
                            # Reduce dimensions by 40%
                            new_size = (int(img.size[0] * 0.6), int(img.size[1] * 0.6))
                            img = img.resize(new_size, Image.Resampling.LANCZOS)
                            # Convert to RGB and reduce colors
                            img = img.convert('RGB')
                            img = img.quantize(colors=128, method=2)
                            img = img.convert('RGB')
                            buffer = io.BytesIO()
                            img.save(buffer, format='JPEG', quality=50, optimize=True)
                            image_part.blob = buffer.getvalue()
                        except Exception as e:
                            print(f"Error in color optimization: {str(e)}")

def apply_lossless_algorithm_1(doc):
    """Structure Optimization: Removes unnecessary elements and optimizes document structure"""
    # Remove comments
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run._element.xpath('.//w:commentReference'):
                run._element.getparent().remove(run._element)
    
    # Remove empty paragraphs
    for paragraph in doc.paragraphs:
        if not paragraph.text and not paragraph._element.xpath('.//w:drawing'):
            paragraph._element.getparent().remove(paragraph._element)
    
    # Remove unused styles
    styles = doc.styles
    for style in styles:
        if not style.element.xpath('.//w:rPr'):
            style.element.getparent().remove(style.element)

class HuffmanNode:
    def __init__(self, char: str, freq: int):
        self.char = char
        self.freq = freq
        self.left = None
        self.right = None

    def __lt__(self, other):
        return self.freq < other.freq

def build_huffman_tree(text: str) -> HuffmanNode:
    """Build a Huffman tree from the input text"""
    # Count character frequencies
    freq = Counter(text)
    
    # Create priority queue
    heap = [HuffmanNode(char, freq) for char, freq in freq.items()]
    heapq.heapify(heap)
    
    # Build tree
    while len(heap) > 1:
        left = heapq.heappop(heap)
        right = heapq.heappop(heap)
        
        internal = HuffmanNode(None, left.freq + right.freq)
        internal.left = left
        internal.right = right
        
        heapq.heappush(heap, internal)
    
    return heap[0]

def build_huffman_codes(root: HuffmanNode, current_code: str = "", codes: dict = None) -> dict:
    """Build Huffman codes from the tree"""
    if codes is None:
        codes = {}
    
    if root is None:
        return codes
    
    if root.char is not None:
        codes[root.char] = current_code if current_code else "0"
    
    build_huffman_codes(root.left, current_code + "0", codes)
    build_huffman_codes(root.right, current_code + "1", codes)
    
    return codes

def huffman_encode(text: str) -> bytes:
    """Huffman coding implementation"""
    # Build Huffman tree and codes
    root = build_huffman_tree(text)
    codes = build_huffman_codes(root)
    
    # Encode text
    encoded = ''.join(codes[char] for char in text)
    
    # Convert to bytes
    padding = 8 - (len(encoded) % 8)
    encoded += '0' * padding
    
    # Convert to bytes
    result = bytearray()
    for i in range(0, len(encoded), 8):
        byte = encoded[i:i+8]
        result.append(int(byte, 2))
    
    # Add padding information
    result.insert(0, padding)
    
    # Add codes dictionary
    codes_str = str(codes)
    result.extend(len(codes_str).to_bytes(4, 'big'))
    result.extend(codes_str.encode())
    
    return bytes(result)

def run_length_encode(text: str) -> bytes:
    """Run-length encoding implementation"""
    if not text:
        return b''
    
    result = bytearray()
    count = 1
    current = text[0]
    
    for char in text[1:]:
        if char == current:
            count += 1
        else:
            result.extend(current.encode())
            result.extend(str(count).encode())
            current = char
            count = 1
    
    # Add the last character and its count
    result.extend(current.encode())
    result.extend(str(count).encode())
    
    return bytes(result)

def lz77_encode(text: str, window_size: int = 4096, lookahead_size: int = 64) -> bytes:
    """LZ77 compression implementation"""
    result = bytearray()
    pos = 0
    
    while pos < len(text):
        # Find the longest match in the window
        best_match = (0, 0)  # (offset, length)
        window_start = max(0, pos - window_size)
        
        for i in range(window_start, pos):
            match_length = 0
            while (pos + match_length < len(text) and 
                   match_length < lookahead_size and 
                   text[i + match_length] == text[pos + match_length]):
                match_length += 1
            
            if match_length > best_match[1]:
                best_match = (pos - i, match_length)
        
        # If we found a match
        if best_match[1] > 2:
            # Encode as (offset, length, next_char)
            result.extend(best_match[0].to_bytes(2, 'big'))
            result.extend(best_match[1].to_bytes(1, 'big'))
            if pos + best_match[1] < len(text):
                result.extend(text[pos + best_match[1]].encode())
            pos += best_match[1] + 1
        else:
            # Encode as (0, 0, char)
            result.extend(b'\x00\x00')
            result.extend(text[pos].encode())
            pos += 1
    
    return bytes(result)

def apply_lossless_algorithm_2(doc):
    """Content Stream Optimization: Compresses text content using our own algorithms"""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text:
                try:
                    # Try different compression algorithms
                    algorithms = [
                        (huffman_encode, 'Huffman'),
                        (run_length_encode, 'RLE'),
                        (lz77_encode, 'LZ77')
                    ]
                    
                    best_compressed = None
                    best_size = float('inf')
                    
                    for algo, name in algorithms:
                        compressed = algo(run.text)
                        if len(compressed) < best_size:
                            best_compressed = compressed
                            best_size = len(compressed)
                    
                    # Store compressed data in a custom property
                    run._element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}compressed', best_compressed.hex())
                    # Clear original text
                    run.text = ''
                except Exception as e:
                    print(f"Error in content compression: {str(e)}")

def apply_lossless_algorithm_3(doc):
    """Object Stream Optimization: Optimizes document structure and uses object streams"""
    # Create a temporary directory
    temp_dir = "temp_docx"
    os.makedirs(temp_dir, exist_ok=True)
    
    try:
        # Save document to temporary location
        temp_path = os.path.join(temp_dir, "temp.docx")
        doc.save(temp_path)
        
        # Extract DOCX contents
        with zipfile.ZipFile(temp_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        # Process document.xml
        doc_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
        if os.path.exists(doc_xml_path):
            tree = ET.parse(doc_xml_path)
            root = tree.getroot()
            
            # Remove unnecessary elements
            for elem in root.findall('.//w:commentReference', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                parent = root.find('.//*[.//w:commentReference]')
                if parent is not None:
                    parent.remove(elem)
            
            # Save modified document.xml
            tree.write(doc_xml_path, encoding='UTF-8', xml_declaration=True)
        
        # Create new DOCX with maximum compression
        with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED, compresslevel=9) as zipf:
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zipf.write(file_path, arcname)
        
        # Reload the document
        doc = Document(temp_path)
        
    finally:
        # Clean up
        shutil.rmtree(temp_dir)

def compress_docx(input_path, output_path, is_lossy=True):
    """Compress a DOCX file using various algorithms"""
    try:
        # Get original size
        original_size = os.path.getsize(input_path)
        
        # Open the document
        doc = Document(input_path)
        
        if is_lossy:
            # Try each lossy algorithm
            algorithms = [
                (apply_lossy_algorithm_1, 'Aggressive Image Compression', 'Reduces image quality and dimensions significantly'),
                (apply_lossy_algorithm_2, 'Smart Image Compression', 'Adaptive compression based on image content'),
                (apply_lossy_algorithm_3, 'Color Space Optimization', 'Focuses on color reduction and DPI adjustment')
            ]
        else:
            # Try each lossless algorithm
            algorithms = [
                (apply_lossless_algorithm_1, 'Structure Optimization', 'Removes unnecessary elements and optimizes document structure'),
                (apply_lossless_algorithm_2, 'Content Stream Optimization', 'Compresses text content using our own algorithms'),
                (apply_lossless_algorithm_3, 'Object Stream Optimization', 'Optimizes document structure and uses object streams')
            ]
        
        best_compressed_size = original_size
        best_algorithm = None
        
        for algo_func, name, description in algorithms:
            try:
                # Create a temporary document
                temp_doc = Document(input_path)
                
                # Apply the algorithm
                algo_func(temp_doc)
                
                # Save to temporary file
                temp_output = f"{output_path}.temp"
                temp_doc.save(temp_output)
                
                # Get compressed size
                compressed_size = os.path.getsize(temp_output)
                
                # If this attempt produced better compression, keep it
                if compressed_size < best_compressed_size:
                    best_compressed_size = compressed_size
                    best_algorithm = (name, description)
                    # Move the temporary file to the final output path
                    shutil.move(temp_output, output_path)
                else:
                    # Remove the temporary file
                    os.remove(temp_output)
                
            except Exception as e:
                print(f"Error in {name}: {str(e)}")
                if os.path.exists(temp_output):
                    os.remove(temp_output)
                continue
        
        # If no compression was successful, return error
        if best_compressed_size >= original_size:
            return {
                'success': False,
                'error': 'Could not achieve compression'
            }
        
        return {
            'success': True,
            'original_size': original_size,
            'compressed_size': best_compressed_size,
            'ratio': original_size / best_compressed_size if best_compressed_size > 0 else 1,
            'algorithm': best_algorithm[0],
            'description': best_algorithm[1]
        }
        
    except Exception as e:
        print(f"Error compressing DOCX: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }
